"""Document loader for PDF and DOCX files.

Parses files into LoadedDocument with per-page text and page number info.
Phase 2 RAG ingestion pipeline — section 1.1 pypdf + python-docx.
"""

from __future__ import annotations

import logging
import pathlib
import re
from collections import Counter
from dataclasses import dataclass, field

import docx
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# Unicode constants for whitespace normalisation
_NBSP = "\u00a0"  # Non-breaking space
_ZWSP = "\u200b"  # Zero-width space
_BOM = "\ufeff"  # Byte-order mark / zero-width no-break space


@dataclass
class PageContent:
    """Single page (or DOCX pseudo-page) of text."""

    text: str
    page_number: int | None  # 1-based for PDF; None for DOCX


@dataclass
class LoadedDocument:
    """Result of loading a PDF or DOCX file."""

    file_name: str
    pages: list[PageContent] = field(default_factory=list)


def _normalize_text(text: str) -> str:
    """Normalise whitespace per PRD section 1.3 cleanup rules."""
    # NBSP -> space, ZWSP / BOM -> remove
    text = text.replace(_NBSP, " ").replace(_ZWSP, "").replace(_BOM, "")
    # Normalise spaces/tabs on each line
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim each line
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()


def _remove_headers_footers(pages: list[str]) -> list[str]:
    """Remove lines that appear in >= 30% of pages (headers/footers).

    Skip heuristic for documents with fewer than 5 pages (PRD section 1.3).
    """
    if len(pages) < 5:
        return pages

    line_counts: Counter[str] = Counter()
    for page_text in pages:
        seen_on_page: set[str] = set()
        for line in page_text.splitlines():
            stripped = line.strip()
            if stripped and stripped not in seen_on_page:
                line_counts[stripped] += 1
                seen_on_page.add(stripped)

    threshold = len(pages) * 0.30
    repeated = {line for line, count in line_counts.items() if count >= threshold}

    if not repeated:
        return pages

    cleaned: list[str] = []
    for page_text in pages:
        new_lines = [line for line in page_text.splitlines() if line.strip() not in repeated]
        cleaned.append("\n".join(new_lines))
    return cleaned


def load_pdf(path: pathlib.Path) -> LoadedDocument:
    """Parse a PDF file into a LoadedDocument.

    Args:
        path: Path to the PDF file.

    Returns:
        LoadedDocument with one PageContent per PDF page.

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    reader = PdfReader(str(path))
    raw_pages: list[str] = []
    for page_obj in reader.pages:
        raw_pages.append(page_obj.extract_text() or "")

    cleaned_pages = _remove_headers_footers(raw_pages)

    pages: list[PageContent] = []
    for i, text in enumerate(cleaned_pages):
        normalised = _normalize_text(text)
        pages.append(PageContent(text=normalised, page_number=i + 1))

    logger.info("Loaded PDF %s: %d pages", path.name, len(pages))
    return LoadedDocument(file_name=path.name, pages=pages)


def load_docx(path: pathlib.Path) -> LoadedDocument:
    """Parse a DOCX file into a LoadedDocument.

    DOCX has no native page boundaries; all paragraphs are stored as a single
    'page' with page_number=None.

    Args:
        path: Path to the DOCX file.

    Returns:
        LoadedDocument with a single PageContent (page_number=None).

    Raises:
        FileNotFoundError: If the file does not exist.
    """
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    normalised = _normalize_text(full_text)

    pages = [PageContent(text=normalised, page_number=None)]
    logger.info("Loaded DOCX %s: %d paragraphs", path.name, len(paragraphs))
    return LoadedDocument(file_name=path.name, pages=pages)
