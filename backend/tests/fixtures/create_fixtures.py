"""Script to generate test fixture files for Phase 2 RAG tests.

Run: cd backend && uv run python tests/fixtures/create_fixtures.py

PDF is generated with reportlab + TTF font (Arial Unicode) so that pypdf
extract_text() returns real Cyrillic characters (not '?????').
"""
from __future__ import annotations

import io
import pathlib

HERE = pathlib.Path(__file__).parent

# Path to a TTF font with full Cyrillic coverage available on macOS.
# Arial Unicode comes pre-installed on macOS and covers Cyrillic + Latin.
_ARIAL_UNICODE_CANDIDATES: list[str] = [
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
]

_DEJAVU_CANDIDATES: list[str] = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/usr/local/share/fonts/DejaVuSans.ttf",
]


def _find_cyrillic_ttf() -> str:
    """Return the first available TTF path that supports Cyrillic."""
    for path in _ARIAL_UNICODE_CANDIDATES + _DEJAVU_CANDIDATES:
        if pathlib.Path(path).exists():
            return path
    raise FileNotFoundError(
        "No Cyrillic-capable TTF font found. "
        "Install Arial Unicode or DejaVuSans.\n"
        f"Searched: {_ARIAL_UNICODE_CANDIDATES + _DEJAVU_CANDIDATES}"
    )


def create_demo_pdf() -> None:
    """Create a realistic Russian work-schedule PDF for RAG ingestion tests.

    Uses reportlab with a TTF font so pypdf extract_text() returns real
    Cyrillic (not '?????').  Content: ~4-6 sections of an internal work
    regulation for migrant workers at ГК ПИК.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font_path = _find_cyrillic_ttf()
    font_name = "CyrillicFont"
    pdfmetrics.registerFont(TTFont(font_name, font_path))

    # Page 1 — timing sections
    page1_lines: list[tuple[float, str]] = [
        (26.0, "Регламент рабочего времени — ГК ПИК"),
        (24.8, ""),
        (24.0, "1. Время рабочей смены"),
        (23.2, "Рабочий день начинается в 8:00 и заканчивается в 17:00."),
        (22.4, "Выход на объект необходимо подтвердить отметкой у прораба."),
        (21.6, ""),
        (20.8, "2. Обеденный перерыв"),
        (20.0, "Обед: с 12:00 до 13:00. Место приёма пищи — столовая на территории."),
        (19.2, "Покидать рабочую зону допускается только в перерыв."),
        (18.4, ""),
        (17.6, "3. Пересменка"),
        (16.8, "Пересменка происходит с 16:45 до 17:00 у бытовки."),
        (16.0, "Работник обязан передать инструменты и доложить о незавершённых задачах."),
        (15.2, "Временный перевод в другое подразделение — с сохранением оплаты по договору."),
        (14.4, ""),
        (13.6, "4. Ночная смена"),
        (12.8, "Ночная смена: с 20:00 до 06:00 следующего дня."),
        (12.0, "Оплачивается с надбавкой 20 % согласно ТК РФ статья 154."),
        (11.2, "Выход на ночную смену только при наличии медицинского допуска."),
    ]

    # Page 2 — illness and HR contact
    page2_lines: list[tuple[float, str]] = [
        (26.0, "5. Действия при болезни или травме"),
        (25.2, "При плохом самочувствии незамедлительно сообщить прорабу."),
        (24.4, "Запрещается выходить на объект с температурой выше 37,5 °C."),
        (23.6, "Больничный лист оформляется через медпункт ГК ПИК или поликлинику."),
        (22.8, "Отсутствие без уважительной причины свыше 4 часов = прогул по ТК РФ."),
        (22.0, ""),
        (21.2, "6. Контакт HR-службы"),
        (20.4, "HR-специалист: Дарья Петрова, тел. +7 (495) 123-45-67."),
        (19.6, "Рабочие часы HR: пн–пт, 09:00–18:00."),
        (18.8, "Email: hr@pik.ru. Вопросы по зарплате — бухгалтерия, доб. 210."),
        (18.0, ""),
        (17.2, "Настоящий регламент обязателен для всех работников ГК ПИК."),
        (16.4, "Нарушение правил влечёт дисциплинарное взыскание согласно ТК РФ."),
        (15.6, "Ознакомление с документом подтверждается подписью при трудоустройстве."),
    ]

    out_path = HERE / "demo_ru.pdf"
    buf = io.BytesIO()

    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    def draw_page(lines: list[tuple[float, str]], is_title: bool = False) -> None:
        for cm_y, text in lines:
            font_size = 14 if (is_title and cm_y >= 25.5) else 11
            c.setFont(font_name, font_size)
            c.drawString(2 * cm, cm_y * cm, text)

    draw_page(page1_lines, is_title=True)
    c.showPage()
    draw_page(page2_lines)
    c.save()

    out_path.write_bytes(buf.getvalue())
    print(f"Created {out_path} ({out_path.stat().st_size} bytes, font={font_path})")

    # Verify Cyrillic extraction immediately
    _verify_pdf_cyrillic(out_path)


def _verify_pdf_cyrillic(pdf_path: pathlib.Path) -> None:
    """Raise AssertionError if pypdf cannot extract Cyrillic from pdf_path."""
    import re

    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    all_text = " ".join(page.extract_text() or "" for page in reader.pages)
    has_cyrillic = bool(re.search(r"[а-яёА-ЯЁ]", all_text))
    question_ratio = all_text.count("?") / max(len(all_text), 1)

    print(f"  Extract sample: {all_text[:120]!r}")
    print(f"  HAS_CYRILLIC={has_cyrillic}  ?-ratio={question_ratio:.3f}")

    if not has_cyrillic:
        raise AssertionError(
            f"PDF {pdf_path} contains no Cyrillic after pypdf extract — "
            "font encoding issue. Check TTF registration."
        )
    if question_ratio > 0.05:
        raise AssertionError(
            f"PDF {pdf_path} has {question_ratio:.1%} '?' characters — "
            "likely encoding corruption."
        )
    print("  OK: Cyrillic text layer verified.")


def create_sample_docx() -> None:
    """Create a minimal DOCX using python-docx (if available) or raw XML."""
    out_path = HERE / "Sample_Policy.docx"
    try:
        import docx  # type: ignore[import]

        doc = docx.Document()
        doc.add_heading("Политика ГК ПИК", level=1)
        doc.add_paragraph(
            "Настоящий документ регламентирует правила работы на объектах ГК ПИК."
        )
        doc.add_paragraph(
            "Раздел 1. Рабочее время. Смена начинается в 08:00 и заканчивается в 17:00."
        )
        doc.add_paragraph(
            "Раздел 2. Безопасность. Ношение каски и спецодежды обязательно на всех объектах."
        )
        doc.add_paragraph(
            "Раздел 3. Проживание. Заселение в общежитие производится по предъявлению паспорта."
        )
        doc.save(str(out_path))
        print(f"Created {out_path} (python-docx)")
    except ImportError:
        # Fallback: write minimal valid OOXML zip
        import zipfile

        content_types = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>"
        )
        rels = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/>'
            "</Relationships>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            "<w:p><w:r><w:t>Политика ГК ПИК</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>Рабочее время: смена начинается в 08:00.</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>Безопасность: каска обязательна на всех объектах.</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>Проживание: заселение по предъявлению паспорта.</w:t></w:r></w:p>"
            "</w:body>"
            "</w:document>"
        )
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types)
            zf.writestr("_rels/.rels", rels)
            zf.writestr("word/document.xml", document_xml)
        out_path.write_bytes(buf.getvalue())
        print(f"Created {out_path} (raw OOXML, {len(buf.getvalue())} bytes)")


if __name__ == "__main__":
    create_demo_pdf()
    create_sample_docx()
    print("Fixtures created successfully.")
